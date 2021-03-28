from docx import Document, opc, oxml
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = oxml.shared.OxmlElement('w:r')
    rPr = oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def get_document(data):
    document = Document()

    title_style = document.styles.add_style('title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.bold = True
    title_font.name = 'Arial'
    title_font.size = Pt(18)
    document.add_paragraph('Notes for ' + data['metadata']['title'], document.styles['title'])

    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    p = document.add_paragraph('', document.styles['Normal'])

    add_hyperlink(p, 'Watch Lecture Video', 'https://youtube.com/watch?v=' + data['metadata']['videoId'])

    for bullet in data['response']:
        if bullet['type'] != 'text':
            continue

        document.add_paragraph(bullet['data'], document.styles['Normal'])
    
    return document